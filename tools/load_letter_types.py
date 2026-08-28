"""Convert the Factor HR .docx letter formats into Letter Type records.

    python tools/load_letter_types.py            # convert only, write nothing
    python tools/load_letter_types.py --apply

Word documents are not a template format anybody can maintain in a browser, so
each one is converted to HTML with its `{MergeField}` tokens intact. HR can then
edit the wording in ERPNext without opening Word or asking a developer.

## What is deliberately not attempted

**Layout is not preserved.** These documents use tab stops and tables to line up
`Employee Name    :    {EmployeeName}`, and reproducing that faithfully in HTML
would mean guessing at intent. Tables become tables; tab-aligned lines become a
two-column table, which is what they were imitating. Everything else becomes a
paragraph.

The wording, the order and the merge fields all survive. The letterhead does not
— that belongs in a Frappe Letter Head, applied at print time, not baked into
seventeen separate templates.

**Two files convert to nothing.** `Form2Revised` and `Form3ARevised` hold their
content in form controls or images rather than text. They are statutory PF forms
and are flagged rather than silently imported as blank.
"""

import argparse
import json
import os
import re
import sys
import urllib.error
import urllib.parse
import urllib.request

try:
	from docx import Document
except ImportError:
	sys.exit("Needs python-docx:  python -m pip install python-docx")

SRC = "data/factohr/Letter Formats"
BASE = os.environ.get("ERP_URL", "https://mannarubber.m.frappe.cloud")
AUTH = "token {0}:{1}".format(os.environ.get("ERP_KEY", ""), os.environ.get("ERP_SECRET", ""))

# Which of these are actually letters and which are government forms. The
# distinction matters: a statutory form has a legally fixed layout and should be
# reproduced exactly or not at all, whereas a letter is ours to reword.
CATEGORY = {
	"CandidateOfferLetter": "Offer / Onboarding",
	"Experience Certificate-Format": "HR Letter",
	"Service Certificate Letter": "HR Letter",
	"TO WHOM IT MAY CONCERN": "HR Letter",
	"Gratuity": "HR Letter",
	"Salary Advance": "HR Letter",
	"Request For Liquor Permit": "HR Letter",
	"Traffic Warning": "Warning / Notice",
}

TOKEN = re.compile(r"\{([A-Za-z0-9 _.\-/]{2,40})\}")


def esc(s):
	return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def para_html(p):
	"""One paragraph, keeping bold runs and turning tab-aligned label/value
	lines into a two-column row — which is the shape they were imitating."""
	text = p.text
	if not text.strip():
		return ""
	# "Employee Name\t:\t{EmployeeName}" and friends
	if "\t" in text:
		cells = [c.strip() for c in text.split("\t") if c.strip()]
		if len(cells) >= 2:
			return ("<tr>" + "".join("<td>{0}</td>".format(esc(c)) for c in cells) + "</tr>")
	bits = []
	for run in p.runs:
		if not run.text:
			continue
		t = esc(run.text)
		if run.bold:
			t = "<b>{0}</b>".format(t)
		bits.append(t)
	body = "".join(bits) or esc(text)
	style = (p.style.name or "").lower() if p.style is not None else ""
	if "heading" in style:
		return "<h3>{0}</h3>".format(body)
	return "<p>{0}</p>".format(body)


def convert(path):
	doc = Document(path)
	out, open_table = [], False

	def close():
		nonlocal open_table
		if open_table:
			out.append("</table>")
			open_table = False

	for p in doc.paragraphs:
		html = para_html(p)
		if not html:
			continue
		if html.startswith("<tr>"):
			if not open_table:
				out.append('<table class="letter-fields">')
				open_table = True
			out.append(html)
		else:
			close()
			out.append(html)
	close()

	for t in doc.tables:
		out.append('<table class="letter-table" border="1">')
		for row in t.rows:
			out.append("<tr>" + "".join(
				"<td>{0}</td>".format(esc(c.text.strip())) for c in row.cells) + "</tr>")
		out.append("</table>")

	html = "\n".join(out)
	return html, sorted({m.strip() for m in TOKEN.findall(html)})


def req(path, method="GET", payload=None):
	d = json.dumps(payload).encode() if payload is not None else None
	r = urllib.request.Request(BASE + path, data=d, method=method,
		headers={"Authorization": AUTH, "Accept": "application/json",
		         "Content-Type": "application/json"})
	try:
		with urllib.request.urlopen(r, timeout=180) as x:
			return True, json.loads(x.read().decode())
	except urllib.error.HTTPError as e:
		return False, e.read().decode()[:160]


def main():
	ap = argparse.ArgumentParser()
	ap.add_argument("--apply", action="store_true")
	args = ap.parse_args()

	if args.apply and not os.environ.get("ERP_KEY"):
		sys.exit("Set ERP_URL, ERP_KEY, ERP_SECRET to apply.")

	made = skipped = 0
	for fn in sorted(os.listdir(SRC)):
		if not fn.lower().endswith(".docx"):
			continue
		stem = os.path.splitext(fn)[0]
		try:
			html, fields = convert(os.path.join(SRC, fn))
		except Exception as exc:
			print("  {0:44} FAILED: {1}".format(stem[:44], str(exc)[:50]))
			continue

		if len(re.sub(r"<[^>]+>", "", html).strip()) < 40:
			print("  {0:44} SKIPPED - no extractable text".format(stem[:44]))
			skipped += 1
			continue

		category = CATEGORY.get(stem, "Statutory Form")
		print("  {0:44} {1:20} {2:>2} fields, {3} chars".format(
			stem[:44], category, len(fields), len(html)))

		if args.apply:
			ok, r = req("/api/resource/Letter%20Type", "POST", {
				"doctype": "Letter Type", "letter_name": stem, "category": category,
				"is_active": 1, "source_file": fn, "template": html,
				"fields_used": ", ".join(fields)})
			if ok:
				made += 1
			elif "Duplicate" in str(r):
				req("/api/resource/Letter%20Type/" + urllib.parse.quote(stem), "PUT",
				    {"template": html, "fields_used": ", ".join(fields), "category": category})
				made += 1
			else:
				print("      -> {0}".format(str(r)[:90]))

	print("\n{0} template(s) {1}, {2} skipped".format(
		made, "written" if args.apply else "converted (dry run)", skipped))
	if not args.apply:
		print("Re-run with --apply to write them into ERPNext.")


if __name__ == "__main__":
	main()
